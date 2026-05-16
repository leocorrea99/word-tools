"""Operações diretas em arquivos .docx usando python-docx."""
from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Tags que podem conter parágrafos aninhados (tabelas, células, SDT, etc.)
_CONTAINER_TAGS = {
    f"{{{_W}}}tbl",
    f"{{{_W}}}tr",
    f"{{{_W}}}tc",
    f"{{{_W}}}sdt",
    f"{{{_W}}}sdtContent",
}
_PARA_TAG = f"{{{_W}}}p"


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _find_caption_style_id(doc: Document) -> str:
    for style in doc.styles:
        sid = (style.style_id or "").lower()
        nm = (style.name or "").lower()
        if sid in ("caption", "legenda") or nm in ("caption", "legenda"):
            return style.style_id
    return "Caption"


def _is_caption_elem(elem, caption_style_id: str) -> bool:
    pPr = elem.find(f"{{{_W}}}pPr")
    if pPr is None:
        return False
    pStyle = pPr.find(f"{{{_W}}}pStyle")
    if pStyle is None:
        return False
    val = pStyle.get(f"{{{_W}}}val", "")
    return val.lower() in {"caption", "legenda", caption_style_id.lower()}


def _has_inline_image(elem) -> bool:
    """Verifica se um <w:p> contém imagem (w:drawing moderno ou w:pict legado)."""
    return bool(
        elem.findall(f".//{{{_W}}}drawing")
        or elem.findall(f".//{{{_W}}}pict")
    )


def _count_images(elem) -> int:
    """Conta imagens lógicas em um <w:p>: 1 por w:r que contém desenho.
    Evita dupla contagem de mc:AlternateContent (Choice + Fallback no mesmo run).
    """
    count = 0
    for r in elem.iter(f"{{{_W}}}r"):
        if r.findall(f".//{{{_W}}}drawing"):
            count += 1
        elif r.findall(f".//{{{_W}}}pict"):
            count += 1
    return count if count > 0 else 1


def _collect_image_paras(container):
    """
    Percorre recursivamente o container e retorna lista de
    (para_elem, parent_elem) para cada parágrafo que contém imagem.
    Cobre parágrafos no corpo E dentro de tabelas.
    """
    result = []
    for child in container:
        if child.tag == _PARA_TAG:
            if _has_inline_image(child):
                result.append((child, container))
        elif child.tag in _CONTAINER_TAGS:
            result.extend(_collect_image_paras(child))
    return result


_ALIGN_OOXML = {"left": "left", "center": "center", "right": "right", "justified": "both"}


def _make_caption_para(
    label_name: str,
    caption_style_id: str,
    alignment: str = "center",
    caption_text: str = "",
):
    """Cria <w:p> com estilo Caption, alinhamento e campo SEQ auto-numerado.
    Se caption_text for fornecido, é inserido após o número (ex: '— Descrição').
    """
    p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), caption_style_id)
    pPr.append(pStyle)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), _ALIGN_OOXML.get(alignment, "center"))
    pPr.append(jc)
    p.append(pPr)

    # Texto do rótulo
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t1.text = label_name + " "
    r1.append(t1)
    p.append(r1)

    # Campo SEQ — início
    r2 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r2.append(fc1)
    p.append(r2)

    # Campo SEQ — instrução
    r3 = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = f" SEQ {label_name} \\* ARABIC "
    r3.append(it)
    p.append(r3)

    # Campo SEQ — separador
    r4 = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r4.append(fc2)
    p.append(r4)

    # Campo SEQ — valor placeholder
    r5 = OxmlElement("w:r")
    t5 = OxmlElement("w:t")
    t5.text = "1"
    r5.append(t5)
    p.append(r5)

    # Campo SEQ — fim
    r6 = OxmlElement("w:r")
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r6.append(fc3)
    p.append(r6)

    # Texto personalizado opcional após o número
    if caption_text:
        r7 = OxmlElement("w:r")
        t7 = OxmlElement("w:t")
        t7.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t7.text = " " + caption_text
        r7.append(t7)
        p.append(r7)

    return p


# ---------------------------------------------------------------------------
# API pública
# ---------------------------------------------------------------------------

def _image_runs(para_elem) -> list:
    """Retorna lista de <w:r> diretos do parágrafo que contêm imagem."""
    result = []
    for child in para_elem:
        if child.tag != f"{{{_W}}}r":
            continue
        if child.findall(f".//{{{_W}}}drawing") or child.findall(f".//{{{_W}}}pict"):
            result.append(child)
    return result


def _split_and_caption(para_elem, parent_elem, idx, label_name, caption_style_id, alignment="center", caption_text="") -> int:
    """
    Divide um parágrafo com N imagens em N pares (parágrafo-imagem + legenda).
    Insere os novos elementos no parent e remove o parágrafo original.
    Retorna o número de legendas inseridas.
    """
    from copy import deepcopy

    runs = _image_runs(para_elem)
    if not runs:
        return 0

    pPr = para_elem.find(f"{{{_W}}}pPr")

    for run in reversed(runs):
        cap = _make_caption_para(label_name, caption_style_id, alignment, caption_text)
        parent_elem.insert(idx + 1, cap)

        new_p = OxmlElement("w:p")
        if pPr is not None:
            new_p.append(deepcopy(pPr))
        new_p.append(deepcopy(run))
        parent_elem.insert(idx + 1, new_p)

    parent_elem.remove(para_elem)
    return len(runs)


def add_captions_to_selected_in_file(
    file_path: str, label_name: str, alignment: str, selected_1based: set, caption_text: str = ""
) -> tuple[int, int]:
    """
    Insere legendas apenas nas imagens cujos índices (1-based, ordem Word) estão em selected_1based.
    Retorna (adicionadas, já_tinham_legenda).
    """
    from copy import deepcopy

    doc = Document(file_path)
    caption_style_id = _find_caption_style_id(doc)

    # Lista plana: uma entrada por imagem, em ordem de documento
    image_paras = _collect_image_paras(doc.element.body)
    flat = []  # (para_elem, parent_elem, local_i_no_para, total_no_para)
    for para_elem, parent_elem in image_paras:
        n = _count_images(para_elem)
        for local_i in range(n):
            flat.append((para_elem, parent_elem, local_i, n))

    # Agrupa por parágrafo para lidar com parágrafos multi-imagem
    para_map: dict = {}
    for idx_1based in sorted(selected_1based):
        i = idx_1based - 1
        if not (0 <= i < len(flat)):
            continue
        para_elem, parent_elem, local_i, n = flat[i]
        pid = id(para_elem)
        if pid not in para_map:
            para_map[pid] = {
                "para": para_elem, "parent": parent_elem,
                "n": n, "selected_locals": set(),
            }
        para_map[pid]["selected_locals"].add(local_i)

    # Processa em ordem reversa de documento para não invalidar índices
    all_para_ids = [id(p) for p, _ in image_paras]
    sorted_pids = [pid for pid in all_para_ids if pid in para_map]

    added = 0
    skipped = 0
    for pid in reversed(sorted_pids):
        info = para_map[pid]
        para_elem = info["para"]
        parent_elem = info["parent"]
        n = info["n"]
        selected_locals = info["selected_locals"]

        siblings = list(parent_elem)
        try:
            idx = siblings.index(para_elem)
        except ValueError:
            continue

        if n == 1:
            already_has = (
                idx + 1 < len(siblings)
                and siblings[idx + 1].tag == _PARA_TAG
                and _is_caption_elem(siblings[idx + 1], caption_style_id)
            )
            if already_has:
                skipped += 1
            else:
                parent_elem.insert(idx + 1, _make_caption_para(label_name, caption_style_id, alignment, caption_text))
                added += 1
        else:
            # Parágrafo multi-imagem: separa e legenda apenas as selecionadas
            runs = _image_runs(para_elem)
            pPr = para_elem.find(f"{{{_W}}}pPr")
            for local_i in reversed(range(n)):
                if local_i >= len(runs):
                    continue
                new_p = OxmlElement("w:p")
                if pPr is not None:
                    new_p.append(deepcopy(pPr))
                new_p.append(deepcopy(runs[local_i]))
                parent_elem.insert(idx + 1, new_p)
                if local_i in selected_locals:
                    parent_elem.insert(idx + 2, _make_caption_para(label_name, caption_style_id, alignment, caption_text))
                    added += 1
            parent_elem.remove(para_elem)

    doc.save(file_path)
    return added, skipped


def add_captions_to_file(file_path: str, label_name: str, alignment: str = "center", caption_text: str = "") -> tuple[int, int]:
    """
    Insere legendas nativas (campo SEQ) em imagens sem legenda.
    - Parágrafos com N>1 imagens: sempre divide e legenda todas (ignora already_has).
    - Parágrafos com 1 imagem: pula se já tiver legenda logo após.
    Recollecta parágrafos a cada modificação para evitar índices desatualizados.
    Retorna (adicionadas, puladas).
    """
    doc = Document(file_path)
    caption_style_id = _find_caption_style_id(doc)

    added = 0
    log = [f"caption_style_id={caption_style_id}"]
    iteration = 0

    while True:
        iteration += 1
        image_paras = _collect_image_paras(doc.element.body)
        log.append(f"\n--- iter {iteration}: {len(image_paras)} parágrafo(s) com imagem ---")
        made_change = False

        for i, (para_elem, parent_elem) in enumerate(image_paras):
            siblings = list(parent_elem)
            try:
                idx = siblings.index(para_elem)
            except ValueError:
                log.append(f"  [{i}] NÃO ENCONTRADO no parent — pulado")
                continue

            n_imgs = _count_images(para_elem)

            # Parágrafo com múltiplas imagens: divide e legenda todas.
            # Não verifica already_has — garante que fotos antes de legenda existente sejam processadas.
            if n_imgs > 1:
                log.append(f"  [{i}] idx={idx} n_imgs={n_imgs} → split+legenda")
                added += _split_and_caption(para_elem, parent_elem, idx, label_name, caption_style_id, alignment, caption_text)
                log.append(f"    → legendas inseridas (total={added})")
                made_change = True
                break

            # Parágrafo com 1 imagem: pula se já tiver legenda
            already_has = (
                idx + 1 < len(siblings)
                and siblings[idx + 1].tag == _PARA_TAG
                and _is_caption_elem(siblings[idx + 1], caption_style_id)
            )
            next_tag = siblings[idx + 1].tag.split("}")[-1] if idx + 1 < len(siblings) else "(fim)"
            next_val = ""
            if idx + 1 < len(siblings):
                ps = siblings[idx + 1].find(f"{{{_W}}}pPr/{{{_W}}}pStyle")
                if ps is not None:
                    next_val = ps.get(f"{{{_W}}}val", "")
            log.append(
                f"  [{i}] idx={idx} próximo=<{next_tag} style='{next_val}'>"
                f" already_has={already_has} n_imgs=1"
            )

            if already_has:
                continue

            parent_elem.insert(idx + 1, _make_caption_para(label_name, caption_style_id, alignment, caption_text))
            added += 1
            log.append(f"    → legenda inserida (total={added})")
            made_change = True
            break

        if not made_change:
            log.append("  → sem mudanças, encerrando")
            break

    with open("/tmp/word_tools_debug.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(log))

    total = sum(_count_images(p) for p, _ in _collect_image_paras(doc.element.body))
    skipped = max(0, total - added)

    doc.save(file_path)
    return added, skipped


# ---------------------------------------------------------------------------
# Configuração de parágrafo das imagens via OOXML
# ---------------------------------------------------------------------------

_CM_TO_TWIPS = 566.9291  # 1 cm em vigésimos de ponto (twips)


def _ensure_pPr(para_elem):
    pPr = para_elem.find(f"{{{_W}}}pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para_elem.insert(0, pPr)
    return pPr


def _ensure_child(parent, local_name):
    tag = f"{{{_W}}}{local_name}"
    child = parent.find(tag)
    if child is None:
        child = OxmlElement(f"w:{local_name}")
        parent.append(child)
    return child


def _image_para_ids_for_indices(image_paras, selected_1based):
    """Retorna set de id(para_elem) para os índices 1-based selecionados."""
    flat = []
    for para_elem, _ in image_paras:
        n = _count_images(para_elem)
        for _ in range(n):
            flat.append(para_elem)
    ids = set()
    for idx in selected_1based:
        i = idx - 1
        if 0 <= i < len(flat):
            ids.add(id(flat[i]))
    return ids


def configure_image_paragraphs_in_file(
    file_path: str,
    alignment: str | None = None,
    left_indent_cm: float | None = None,
    right_indent_cm: float | None = None,
    space_before_cm: float | None = None,
    space_after_cm: float | None = None,
    line_spacing: str | None = None,
    selected_1based: set | None = None,
) -> int:
    """
    Configura o formato do parágrafo das imagens via OOXML.
    selected_1based=None → todas; caso contrário, apenas os índices indicados.
    Retorna número de parágrafos modificados.
    """
    doc = Document(file_path)
    image_paras = _collect_image_paras(doc.element.body)

    target_ids = None
    if selected_1based is not None:
        target_ids = _image_para_ids_for_indices(image_paras, selected_1based)

    processed = set()
    count = 0

    for para_elem, _ in image_paras:
        pid = id(para_elem)
        if pid in processed:
            continue
        if target_ids is not None and pid not in target_ids:
            continue
        processed.add(pid)

        pPr = _ensure_pPr(para_elem)

        if alignment and alignment in _ALIGN_OOXML:
            jc = _ensure_child(pPr, "jc")
            jc.set(qn("w:val"), _ALIGN_OOXML[alignment])

        if left_indent_cm is not None or right_indent_cm is not None:
            ind = _ensure_child(pPr, "ind")
            if left_indent_cm is not None:
                ind.set(qn("w:left"), str(int(round(left_indent_cm * _CM_TO_TWIPS))))
            if right_indent_cm is not None:
                ind.set(qn("w:right"), str(int(round(right_indent_cm * _CM_TO_TWIPS))))

        if space_before_cm is not None or space_after_cm is not None or line_spacing:
            sp = _ensure_child(pPr, "spacing")
            if space_before_cm is not None:
                sp.set(qn("w:before"), str(int(round(space_before_cm * _CM_TO_TWIPS))))
            if space_after_cm is not None:
                sp.set(qn("w:after"), str(int(round(space_after_cm * _CM_TO_TWIPS))))
            if line_spacing == "single":
                sp.set(qn("w:line"), "240")
                sp.set(qn("w:lineRule"), "auto")
            elif line_spacing == "1.5":
                sp.set(qn("w:line"), "360")
                sp.set(qn("w:lineRule"), "auto")

        count += 1

    doc.save(file_path)
    return count


def set_image_first_line_indent_in_file(
    file_path: str,
    indent_cm: float = 0.0,
    selected_1based: set | None = None,
) -> int:
    """
    Define recuo de primeira linha dos parágrafos de imagens via OOXML.
    indent_cm=0 remove o recuo. Retorna número de parágrafos modificados.
    """
    doc = Document(file_path)
    image_paras = _collect_image_paras(doc.element.body)

    target_ids = None
    if selected_1based is not None:
        target_ids = _image_para_ids_for_indices(image_paras, selected_1based)

    indent_twips = str(int(round(indent_cm * _CM_TO_TWIPS)))

    processed = set()
    count = 0

    for para_elem, _ in image_paras:
        pid = id(para_elem)
        if pid in processed:
            continue
        if target_ids is not None and pid not in target_ids:
            continue
        processed.add(pid)

        pPr = _ensure_pPr(para_elem)
        ind = _ensure_child(pPr, "ind")
        ind.set(qn("w:firstLine"), indent_twips)
        count += 1

    doc.save(file_path)
    return count
